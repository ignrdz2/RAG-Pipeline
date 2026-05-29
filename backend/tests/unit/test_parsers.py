from pathlib import Path

import pytest


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

def _make_pdf(tmp_path: Path, text: str | None = "Hello from page one") -> Path:
    import fitz

    doc = fitz.open()
    page = doc.new_page()
    if text:
        page.insert_text((50, 72), text)
    pdf_path = tmp_path / "test.pdf"
    doc.save(str(pdf_path))
    doc.close()
    return pdf_path


def test_pdf_parser(tmp_path: Path) -> None:
    from app.services.ingestion.parsers.pdf import parse

    pdf_path = _make_pdf(tmp_path, "Hello from page one")
    result = parse(pdf_path)

    assert len(result) == 1
    page_num, text = result[0]
    assert page_num == 1
    assert "Hello from page one" in text


def test_pdf_multipagina(tmp_path: Path) -> None:
    import fitz
    from app.services.ingestion.parsers.pdf import parse

    doc = fitz.open()
    for i, content in enumerate(["Página uno", "Página dos", "Página tres"]):
        page = doc.new_page()
        page.insert_text((50, 72), content)
    pdf_path = tmp_path / "multi.pdf"
    doc.save(str(pdf_path))
    doc.close()

    result = parse(pdf_path)

    assert len(result) == 3
    assert result[0][0] == 1
    assert result[1][0] == 2
    assert result[2][0] == 3
    assert "Página uno" in result[0][1]


def test_pdf_sin_texto(tmp_path: Path) -> None:
    from app.services.ingestion.parsers.pdf import parse

    # Página en blanco — sin texto insertado
    pdf_path = _make_pdf(tmp_path, text=None)

    with pytest.raises(ValueError, match="PDF sin texto extraíble"):
        parse(pdf_path)


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

def _make_docx(tmp_path: Path, paragraphs: list[str]) -> Path:
    from docx import Document

    doc = Document()
    # python-docx añade un párrafo vacío por defecto; lo limpiamos
    for para in doc.paragraphs:
        p = para._element
        p.getparent().remove(p)
    for text in paragraphs:
        doc.add_paragraph(text)
    docx_path = tmp_path / "test.docx"
    doc.save(str(docx_path))
    return docx_path


def test_docx_parser(tmp_path: Path) -> None:
    from app.services.ingestion.parsers.docx import parse

    docx_path = _make_docx(tmp_path, ["Primer párrafo", "Segundo párrafo"])
    result = parse(docx_path)

    assert len(result) == 1
    page_num, text = result[0]
    assert page_num is None
    assert "Primer párrafo" in text
    assert "Segundo párrafo" in text
    # Los párrafos se unen con doble salto de línea
    assert "\n\n" in text


def test_docx_parrafos_vacios_ignorados(tmp_path: Path) -> None:
    from app.services.ingestion.parsers.docx import parse

    docx_path = _make_docx(tmp_path, ["Contenido real", "", "  ", "Más contenido"])
    result = parse(docx_path)

    _, text = result[0]
    assert "Contenido real" in text
    assert "Más contenido" in text
    # Los párrafos vacíos no generan saltos extra al inicio/fin
    assert not text.startswith("\n")
    assert not text.endswith("\n")


# ---------------------------------------------------------------------------
# TXT
# ---------------------------------------------------------------------------

def test_txt_parser(tmp_path: Path) -> None:
    from app.services.ingestion.parsers.txt import parse

    txt_path = tmp_path / "test.txt"
    txt_path.write_text("Hello, world!\nSegunda línea.", encoding="utf-8")

    result = parse(txt_path)

    assert len(result) == 1
    page_num, text = result[0]
    assert page_num is None
    assert "Hello, world!" in text
    assert "Segunda línea." in text


def test_txt_parser_fallback_latin1(tmp_path: Path) -> None:
    from app.services.ingestion.parsers.txt import parse

    txt_path = tmp_path / "latin.txt"
    txt_path.write_bytes("Café con leche\n".encode("latin-1"))

    result = parse(txt_path)

    assert len(result) == 1
    assert result[0][0] is None
    assert "Caf" in result[0][1]  # 'é' puede variar según decodificación


# ---------------------------------------------------------------------------
# Orquestador
# ---------------------------------------------------------------------------

def test_parser_orchestrator_delega_pdf(tmp_path: Path) -> None:
    from app.services.ingestion.parser import parse

    pdf_path = _make_pdf(tmp_path, "Contenido PDF")
    result = parse(pdf_path, "pdf")

    assert result[0][0] == 1  # tiene número de página
    assert "Contenido PDF" in result[0][1]


def test_parser_orchestrator_delega_txt(tmp_path: Path) -> None:
    from app.services.ingestion.parser import parse

    txt_path = tmp_path / "file.txt"
    txt_path.write_text("Texto plano", encoding="utf-8")
    result = parse(txt_path, "txt")

    assert result[0][0] is None
    assert "Texto plano" in result[0][1]


def test_parser_orchestrator_tipo_no_soportado(tmp_path: Path) -> None:
    from app.services.ingestion.parser import parse

    path = tmp_path / "file.xyz"
    path.write_text("content")

    with pytest.raises(ValueError, match="no soportado"):
        parse(path, "xyz")
